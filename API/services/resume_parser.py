"""
Resume Parser Service
Extracts text from PDF, DOCX, and TXT files.
Uses the parsed text to build a structured resume object for screening.
"""

import re
from io import BytesIO
from pathlib import Path

import pdfplumber
from docx import Document


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """
    Extract text from file bytes (no temp file - avoids Windows file lock).
    Supports: PDF, DOCX, TXT
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        bio = BytesIO(data)
        bio.seek(0)
        text_parts = []
        with pdfplumber.open(bio) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    elif suffix == ".docx":
        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    elif suffix == ".txt":
        return data.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {suffix}")


def extract_text_from_file(file_path: str) -> str:
    """
    Extract raw text from a resume file.
    Supports: PDF, DOCX, TXT
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Extracted text as a single string
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    if suffix == ".pdf":
        return _extract_from_pdf(file_path)
    elif suffix == ".docx":
        return _extract_from_docx(file_path)
    elif suffix == ".txt":
        return _extract_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _extract_from_txt(file_path: str) -> str:
    """Extract text from plain TXT file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def parse_resume_text(text: str) -> dict:
    """
    Parse raw resume text into a structured format for screening.
    Extracts: skills, experience years, education, and full text.
    
    Args:
        text: Raw resume text
        
    Returns:
        Dictionary with parsed resume data
    """
    text = text.strip()
    if not text:
        return {"full_text": "", "skills": [], "experience_years": 0, "education": []}
    
    # Extract skills (common tech/business terms - simplified keyword extraction)
    skills = _extract_skills(text)
    
    # Extract years of experience
    experience_years = _extract_experience_years(text)
    
    # Extract education keywords
    education = _extract_education(text)
    
    return {
        "full_text": text,
        "skills": skills,
        "experience_years": experience_years,
        "education": education,
    }


def _extract_skills(text: str) -> list:
    """
    Extract skill keywords from resume text.
    Uses common technical and business skills as a base list.
    """
    # Common skills to look for (extend this list as needed)
    skill_keywords = [
        "python", "java", "javascript", "sql", "excel", "power bi", "tableau",
        "data analysis", "machine learning", "statistics", "data visualization",
        "etl", "databases", "mongodb", "mysql", "postgresql", "aws", "azure",
        "git", "agile", "scrum", "project management", "communication",
        "problem solving", "data mining", "reporting", "business intelligence",
        "r", "pandas", "numpy", "spss", "sas", "api", "rest", "etl"
    ]
    
    text_lower = text.lower()
    found_skills = []
    for skill in skill_keywords:
        if skill in text_lower:
            found_skills.append(skill)
    
    return list(set(found_skills))


def _extract_experience_years(text: str) -> int:
    """
    Try to extract years of experience from resume.
    Looks for patterns like "5 years", "3+ years", "10 years of experience".
    """
    patterns = [
        r"(\d+)\s*\+\s*years?\s*(?:of\s+)?(?:experience|exp)",
        r"(\d+)\s*years?\s*(?:of\s+)?(?:experience|exp)",
        r"experience[:\s]+(\d+)\s*years?",
        r"(\d+)\s*years?\s+experience",
    ]
    
    text_lower = text.lower()
    for pattern in patterns:
        match = re.search(pattern, text_lower, re.IGNORECASE)
        if match:
            return min(int(match.group(1)), 30)  # Cap at 30 for sanity
    
    return 0


def _extract_education(text: str) -> list:
    """Extract education-related terms (degrees, institutions)."""
    education_keywords = [
        "bachelor", "master", "phd", "mba", "btech", "mtech", "bsc", "msc",
        "degree", "graduation", "computer science", "engineering", "mathematics",
        "statistics", "business administration"
    ]
    
    text_lower = text.lower()
    found = []
    for keyword in education_keywords:
        if keyword in text_lower:
            found.append(keyword)
    
    return list(set(found))
